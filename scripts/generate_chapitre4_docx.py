from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "chapitre4_revise.md"
OUTPUT = ROOT / "docs" / "Chapitre_4_revise_3mmalak.docx"
PDF_OUTPUT = ROOT / "docs" / "Chapitre_4_revise_3mmalak.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 22, RGBColor(15, 118, 110)),
        ("Heading 1", 16, RGBColor(31, 41, 55)),
        ("Heading 2", 13, RGBColor(15, 118, 110)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Body Small" not in styles:
        body_small = styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        body_small.base_style = styles["Normal"]
        body_small.font.name = "Arial"
        body_small.font.size = Pt(10)
        body_small.font.color.rgb = RGBColor(90, 90, 90)


def add_header(document):
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Plateforme 3mmalak")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(15, 118, 110)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Chapitre 4 - Realisation")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)


def add_cover_box(document):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    set_cell_shading(cell, "EAF7F5")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version revisee et alignee avec la plateforme implementee")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 118, 110)

    p2 = cell.add_paragraph(
        "Ce document remplace le contenu generique du chapitre 4 par une presentation conforme au projet actuel : roles utilisateur, gestion des taches, reservations, administration, cartes, notifications et verification documentaire."
    )
    p2.style = "Body Small"
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_markdown_lines(text):
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            yield ("blank", "")
        elif line.startswith("# "):
            yield ("title", line[2:].strip())
        elif line.startswith("## "):
            yield ("h1", line[3:].strip())
        elif line.startswith("### "):
            yield ("h2", line[4:].strip())
        elif line.startswith("- "):
            yield ("bullet", line[2:].strip())
        else:
            yield ("p", line)


def build_pdf():
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0F766E"),
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    subtitle_style = ParagraphStyle(
        "CustomSubtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#6B7280"),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h1_style = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#1F2937"),
        spaceBefore=10,
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#0F766E"),
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )

    story = []
    bullet_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        if not bullet_buffer:
            return
        items = [
            ListItem(Paragraph(item.replace("&", "&amp;"), body_style), leftIndent=8)
            for item in bullet_buffer
        ]
        story.append(
            ListFlowable(
                items,
                bulletType="bullet",
                start="circle",
                leftIndent=14,
            )
        )
        story.append(Spacer(1, 0.15 * cm))
        bullet_buffer = []

    for kind, content in parse_markdown_lines(SOURCE.read_text(encoding="utf-8")):
        safe = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind == "title":
            box = Table(
                [[Paragraph(safe, title_style)]],
                colWidths=[16.2 * cm],
            )
            box.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF7F5")),
                        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#B7E4DD")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 16),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 16),
                        ("TOPPADDING", (0, 0), (-1, -1), 12),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                    ]
                )
            )
            story.append(box)
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("Memoire de fin d'etudes - version adaptee au projet web reel", subtitle_style))
            story.append(Spacer(1, 0.35 * cm))
        elif kind == "h1":
            flush_bullets()
            story.append(Paragraph(safe, h1_style))
        elif kind == "h2":
            flush_bullets()
            story.append(Paragraph(safe, h2_style))
        elif kind == "bullet":
            bullet_buffer.append(safe)
        elif kind == "p":
            flush_bullets()
            story.append(Paragraph(safe, body_style))
        else:
            flush_bullets()
            story.append(Spacer(1, 0.08 * cm))

    flush_bullets()

    doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2.2 * cm,
        title="Chapitre 4 - Realisation de la plateforme 3mmalak",
    )
    doc.build(story)


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    configure_styles(document)
    add_header(document)

    for kind, content in parse_markdown_lines(SOURCE.read_text(encoding="utf-8")):
        if kind == "title":
            p = document.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(content)
            subtitle = document.add_paragraph(
                "Memoire de fin d'etudes - version adaptee au projet web reel"
            )
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.style = "Body Small"
            add_cover_box(document)
            document.add_paragraph("")
        elif kind == "h1":
            document.add_paragraph(content, style="Heading 1")
        elif kind == "h2":
            document.add_paragraph(content, style="Heading 2")
        elif kind == "bullet":
            document.add_paragraph(content, style="List Bullet")
        elif kind == "p":
            document.add_paragraph(content, style="Normal")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    build_pdf()
    print(OUTPUT)
    print(PDF_OUTPUT)


if __name__ == "__main__":
    build_document()
